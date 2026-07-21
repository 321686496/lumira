# -*- coding: utf-8 -*-
"""
将「如画 Lumira 摄影辅助应用 V1.0.0」软著 Markdown 文档转换为符合规范的 docx 文档。

软著提交材料文档规范：
- 纸张：A4（210mm × 297mm）
- 页边距：上下 2.54cm，左 3cm（装订），右 2.54cm
- 页眉：软件名称 + 版本号
- 页脚：页码居中（第 X 页）
- 正文：宋体 小四（12pt），1.5 倍行距，首行缩进 2 字符
- 一级标题（文档主标题）：黑体 二号（22pt）居中
- 二级标题（一、二、三…）：黑体 三号（16pt）
- 三级标题（1.1、2.1…）：黑体 四号（14pt）
- 四级标题（2.1.1…）：黑体 小四（12pt）
- 表格：宋体 五号（10.5pt）
- 代码块：等线 五号（10.5pt），无缩进，保留行号
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ====================== 全局配置 ======================

SOFTWARE_NAME = "如画 Lumira 摄影辅助应用 V1.0.0"
OUTPUT_DIR = r"d:\app\projects\photo_post\软件著作权\docx"
INPUT_DIR = r"d:\app\projects\photo_post\软件著作权"

# 字体配置
FONT_SONG = "宋体"          # 正文中文字体
FONT_HEI = "黑体"           # 标题中文字体
FONT_ASCII = "Times New Roman"  # 正文西文字体
FONT_CODE = "Consolas"      # 代码字体

# 字号配置
FS_DOC_TITLE = Pt(22)       # 二号 - 文档主标题
FS_H1 = Pt(16)              # 三号 - 一级标题
FS_H2 = Pt(14)              # 四号 - 二级标题
FS_H3 = Pt(12)              # 小四 - 三级标题
FS_BODY = Pt(12)            # 小四 - 正文
FS_TABLE = Pt(10.5)         # 五号 - 表格
FS_CODE = Pt(10.5)          # 五号 - 代码

# 颜色
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_GRAY = RGBColor(0x80, 0x80, 0x80)


# ====================== 通用工具函数 ======================

def set_run_font(run, cn_font=FONT_SONG, ascii_font=FONT_ASCII,
                 size=FS_BODY, bold=False, color=COLOR_BLACK):
    """统一设置 run 的中英文字体、字号、粗体、颜色。"""
    run.font.name = ascii_font
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    # 中文字体通过 w:rFonts 的 w:eastAsia 设置
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), ascii_font)


def set_paragraph_format(paragraph, alignment=None, line_spacing=1.5,
                         first_line_indent=None, space_before=Pt(0),
                         space_after=Pt(0), left_indent=None):
    """统一设置段落格式。"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if left_indent is not None:
        pf.left_indent = left_indent


def set_page_layout(doc):
    """A4 纸张 + 软著标准页边距（左 3cm 装订）。"""
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


def add_page_header(section, text=SOFTWARE_NAME):
    """添加页眉：软件名称 + 版本号（居中，宋体五号）。"""
    header = section.header
    header.is_linked_to_previous = False
    # 清空已有段落
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, cn_font=FONT_SONG, ascii_font=FONT_ASCII,
                 size=Pt(10.5), bold=False)
    # 添加底部边框（页眉横线）
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_number_field(paragraph):
    """在页脚段落中插入「第 X 页」页码域（居中）。"""
    # 「第 」前缀
    run_pre = paragraph.add_run("第 ")
    set_run_font(run_pre, size=Pt(10.5))

    # PAGE 域
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE \\* MERGEFORMAT '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    fldText = OxmlElement('w:t')
    fldText.text = '1'
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run_field = paragraph.add_run()
    run_field._r.append(fldChar_begin)
    run_field._r.append(instrText)
    run_field._r.append(fldChar_sep)
    run_field._r.append(fldText)
    run_field._r.append(fldChar_end)
    set_run_font(run_field, size=Pt(10.5))

    # 「 页」后缀
    run_post = paragraph.add_run(" 页")
    set_run_font(run_post, size=Pt(10.5))


def add_page_footer(section):
    """添加页脚：居中页码。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(p)


def setup_document_styles(doc):
    """配置文档默认样式：宋体小四 + 1.5 倍行距。"""
    # Normal 样式
    normal = doc.styles['Normal']
    normal.font.name = FONT_ASCII
    normal.font.size = FS_BODY
    rFonts = normal.element.rPr.rFonts if normal.element.rPr is not None and normal.element.rPr.rFonts is not None else None
    if rFonts is None:
        rPr = normal.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_ASCII)
    rFonts.set(qn('w:hAnsi'), FONT_ASCII)
    rFonts.set(qn('w:eastAsia'), FONT_SONG)
    rFonts.set(qn('w:cs'), FONT_ASCII)

    # 段落格式：1.5 倍行距
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ====================== Markdown 解析与渲染 ======================

class MarkdownParser:
    """简易 Markdown 解析器，识别软著文档中使用的语法。"""

    def __init__(self, text):
        self.lines = text.splitlines()
        self.i = 0
        self.n = len(self.lines)

    def parse(self):
        """返回一个 block 列表，每个 block 为 dict:
        {type: 'h1'|'h2'|'h3'|'h4'|'p'|'table'|'code'|'list_ordered'|'list_unordered'|'hr'|'quote',
         ...}
        """
        blocks = []
        while self.i < self.n:
            line = self.lines[self.i]

            # 跳过 HTML 注释（页眉/页码占位）
            if line.strip().startswith('<!--'):
                # 跳过单行 HTML 注释
                self.i += 1
                continue

            # 空行
            if not line.strip():
                self.i += 1
                continue

            # 水平分隔线
            if line.strip() in ('---', '***', '___'):
                blocks.append({'type': 'hr'})
                self.i += 1
                continue

            # 标题
            m = re.match(r'^(#{1,4})\s+(.*)$', line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                # 去掉 Markdown 强调符号
                text = self._strip_inline_md(text)
                blocks.append({'type': f'h{level}', 'text': text})
                self.i += 1
                continue

            # 引用块（> ...）
            if line.lstrip().startswith('>'):
                quote_lines = []
                while self.i < self.n and self.lines[self.i].lstrip().startswith('>'):
                    content = re.sub(r'^\s*>\s?', '', self.lines[self.i])
                    quote_lines.append(content)
                    self.i += 1
                # 对引用内容同样去除行内 Markdown 强调符号
                quote_lines = [self._strip_inline_md(l) for l in quote_lines]
                blocks.append({'type': 'quote', 'lines': quote_lines})
                continue

            # 代码块
            if line.lstrip().startswith('```'):
                # 取得缩进
                indent = len(line) - len(line.lstrip())
                # 代码块内的所有行（直到闭合 ```）
                code_lines = []
                self.i += 1
                while self.i < self.n:
                    cur = self.lines[self.i]
                    if cur.lstrip().startswith('```'):
                        self.i += 1
                        break
                    # 去除与起始行相同的缩进
                    if indent > 0 and cur[:indent].isspace():
                        code_lines.append(cur[indent:])
                    else:
                        code_lines.append(cur)
                    self.i += 1
                blocks.append({'type': 'code', 'lines': code_lines})
                continue

            # 表格（以 | 开头）
            if line.lstrip().startswith('|'):
                table_lines = []
                while self.i < self.n and self.lines[self.i].lstrip().startswith('|'):
                    table_lines.append(self.lines[self.i].strip())
                    self.i += 1
                blocks.append({'type': 'table', 'lines': table_lines})
                continue

            # 有序列表
            if re.match(r'^\s*\d+\.\s+', line):
                list_items = []
                base_indent = len(line) - len(line.lstrip())
                while self.i < self.n:
                    cur = self.lines[self.i]
                    if not cur.strip():
                        self.i += 1
                        # 允许列表项之间的空行，检查下一行是否仍是列表
                        if self.i < self.n and re.match(r'^\s*\d+\.\s+', self.lines[self.i]):
                            continue
                        else:
                            break
                    if re.match(r'^\s*\d+\.\s+', cur):
                        item_text = re.sub(r'^\s*\d+\.\s+', '', cur)
                        list_items.append(self._strip_inline_md(item_text))
                        self.i += 1
                    else:
                        # 列表项的续行（缩进更深）
                        if list_items:
                            list_items[-1] += ' ' + cur.strip()
                            self.i += 1
                        else:
                            break
                blocks.append({'type': 'list_ordered', 'items': list_items})
                continue

            # 无序列表
            if re.match(r'^\s*[-*+]\s+', line) and not line.lstrip().startswith('***'):
                list_items = []
                while self.i < self.n:
                    cur = self.lines[self.i]
                    if not cur.strip():
                        self.i += 1
                        if self.i < self.n and re.match(r'^\s*[-*+]\s+', self.lines[self.i]):
                            continue
                        else:
                            break
                    m = re.match(r'^\s*[-*+]\s+(.*)$', cur)
                    if m:
                        list_items.append(self._strip_inline_md(m.group(1)))
                        self.i += 1
                    else:
                        if list_items:
                            list_items[-1] += ' ' + cur.strip()
                            self.i += 1
                        else:
                            break
                blocks.append({'type': 'list_unordered', 'items': list_items})
                continue

            # 普通段落
            para_lines = [line.strip()]
            self.i += 1
            while self.i < self.n:
                cur = self.lines[self.i]
                if not cur.strip():
                    break
                # 以下情况结束段落
                if cur.lstrip().startswith(('#', '>', '```', '|', '---', '***', '___')):
                    break
                if re.match(r'^\s*\d+\.\s+', cur) or re.match(r'^\s*[-*+]\s+', cur):
                    break
                para_lines.append(cur.strip())
                self.i += 1
            para_text = ' '.join(para_lines)
            blocks.append({'type': 'p', 'text': para_text})
        return blocks

    @staticmethod
    def _strip_inline_md(text):
        """去除 Markdown 行内强调符号（**bold**, *italic*, `code`），保留纯文本。"""
        # 代码 `code`
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # 粗体 **bold**
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        # 斜体 *italic*
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        # 链接 [text](url)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text


# ====================== Markdown 渲染到 docx ======================

def render_block(doc, block, code_page_no=None):
    """将单个 block 渲染为 docx 内容。"""
    bt = block['type']

    if bt == 'h1':
        # 文档主标题：黑体二号居中
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.5, first_line_indent=None,
                             space_before=Pt(12), space_after=Pt(18))
        run = p.add_run(block['text'])
        set_run_font(run, cn_font=FONT_HEI, ascii_font=FONT_ASCII,
                     size=FS_DOC_TITLE, bold=True)

    elif bt == 'h2':
        # 一、二、三...：黑体三号
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, first_line_indent=None,
                             space_before=Pt(18), space_after=Pt(8))
        run = p.add_run(block['text'])
        set_run_font(run, cn_font=FONT_HEI, ascii_font=FONT_ASCII,
                     size=FS_H1, bold=True)

    elif bt == 'h3':
        # 1.1、2.1...：黑体四号
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, first_line_indent=None,
                             space_before=Pt(10), space_after=Pt(6))
        run = p.add_run(block['text'])
        set_run_font(run, cn_font=FONT_HEI, ascii_font=FONT_ASCII,
                     size=FS_H2, bold=True)

    elif bt == 'h4':
        # 2.1.1...：黑体小四
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             line_spacing=1.5, first_line_indent=None,
                             space_before=Pt(8), space_after=Pt(4))
        run = p.add_run(block['text'])
        set_run_font(run, cn_font=FONT_HEI, ascii_font=FONT_ASCII,
                     size=FS_H3, bold=True)

    elif bt == 'p':
        # 正文：宋体小四 + 首行缩进 2 字符
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             line_spacing=1.5,
                             first_line_indent=Pt(FS_BODY.pt * 2),
                             space_before=Pt(0), space_after=Pt(4))
        run = p.add_run(block['text'])
        set_run_font(run, size=FS_BODY)

    elif bt == 'quote':
        # 引用块：宋体小四斜体，左缩进
        for i, line in enumerate(block['lines']):
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.5,
                                 first_line_indent=None,
                                 left_indent=Cm(1),
                                 space_before=Pt(2), space_after=Pt(2))
            run = p.add_run(line)
            set_run_font(run, size=FS_BODY, color=COLOR_GRAY)

    elif bt == 'list_ordered':
        for idx, item in enumerate(block['items'], 1):
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                 line_spacing=1.5,
                                 left_indent=Cm(0.74),
                                 first_line_indent=Pt(-FS_BODY.pt * 2),
                                 space_before=Pt(0), space_after=Pt(2))
            run = p.add_run(f"{idx}. {item}")
            set_run_font(run, size=FS_BODY)

    elif bt == 'list_unordered':
        for item in block['items']:
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                 line_spacing=1.5,
                                 left_indent=Cm(0.74),
                                 first_line_indent=Pt(-FS_BODY.pt * 2),
                                 space_before=Pt(0), space_after=Pt(2))
            run = p.add_run(f"• {item}")
            set_run_font(run, size=FS_BODY)

    elif bt == 'table':
        render_table(doc, block['lines'])

    elif bt == 'code':
        render_code_block(doc, block['lines'])

    elif bt == 'hr':
        # 水平分隔线：插入底部边框的空段落
        p = doc.add_paragraph()
        set_paragraph_format(p, line_spacing=1.5, space_before=Pt(6),
                             space_after=Pt(6))
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)


def render_table(doc, table_lines):
    """渲染 Markdown 表格为 docx 表格。"""
    # 解析表格
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # 分隔行（|---|---|）跳过
        if re.match(r'^\|[\s:|-]+\|$', line):
            continue
        # 拆分单元格
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    # 第一行为表头
    header = rows[0]
    body_rows = rows[1:]
    n_cols = len(header)

    # 创建表格
    table = doc.add_table(rows=1 + len(body_rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置表格样式（网格）
    table.style = 'Table Grid'

    # 表头
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(header):
        if i >= n_cols:
            break
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 清空默认段落
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                             line_spacing=1.2, first_line_indent=None,
                             space_before=Pt(0), space_after=Pt(0))
        # 加粗表头
        run = p.add_run(MarkdownParser._strip_inline_md(cell_text))
        set_run_font(run, cn_font=FONT_SONG, ascii_font=FONT_ASCII,
                     size=FS_TABLE, bold=True)
        # 表头背景色（浅灰）
        _set_cell_shading(cell, 'D9D9D9')

    # 数据行
    for r_idx, row_data in enumerate(body_rows):
        row = table.rows[r_idx + 1]
        for c_idx in range(n_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ''
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.2, first_line_indent=None,
                                 space_before=Pt(0), space_after=Pt(0))
            run = p.add_run(MarkdownParser._strip_inline_md(cell_text))
            set_run_font(run, size=FS_TABLE)


def _set_cell_shading(cell, color_hex):
    """给单元格设置背景色。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def render_code_block(doc, code_lines):
    """渲染代码块：等线五号，无缩进，保留行号。"""
    # 软著规范：源代码每页 50 行。若代码块超过 50 行，自动分页。
    PAGE_LINES = 50
    # 判断是否带有行号（行首形如 "   1  code"）
    has_line_numbers = any(re.match(r'^\s*\d+\s{2,}', line) for line in code_lines[:5])

    chunks = [code_lines[i:i + PAGE_LINES]
              for i in range(0, len(code_lines), PAGE_LINES)] if has_line_numbers else [code_lines]

    for chunk_idx, chunk in enumerate(chunks):
        if chunk_idx > 0:
            # 分页符
            doc.add_page_break()
        for line in chunk:
            p = doc.add_paragraph()
            set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 line_spacing=1.0,
                                 first_line_indent=None,
                                 left_indent=None,
                                 space_before=Pt(0), space_after=Pt(0))
            # 使用等宽字体显示代码
            run = p.add_run(line if line else ' ')
            set_run_font(run, cn_font=FONT_SONG, ascii_font=FONT_CODE,
                         size=FS_CODE, color=COLOR_BLACK)


# ====================== 文档生成主流程 ======================

def convert_md_to_docx(md_path, docx_path, doc_title=None):
    """将单个 Markdown 文件转换为 docx。"""
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    setup_document_styles(doc)
    set_page_layout(doc)

    # 页眉页脚
    for section in doc.sections:
        add_page_header(section)
        add_page_footer(section)

    # 解析 Markdown
    parser = MarkdownParser(text)
    blocks = parser.parse()

    # 如果指定了文档标题，在文档开头插入封面标题（不进页眉页脚的设置）
    if doc_title:
        # 文档主标题作为首个 H1（实际上 h1 也已是主标题）
        pass

    # 渲染所有 block
    for block in blocks:
        render_block(doc, block)

    # 保存
    doc.save(docx_path)
    return docx_path


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    files = [
        ("01-软件著作权登记申请表.md", "01-软件著作权登记申请表.docx"),
        ("02-软件功能说明书.md", "02-软件功能说明书.docx"),
        ("03-软件用户手册.md", "03-软件用户手册.docx"),
        ("04-源代码说明文档.md", "04-源代码说明文档.docx"),
    ]

    print(f"输出目录：{OUTPUT_DIR}")
    print(f"软件名称：{SOFTWARE_NAME}")
    print("-" * 60)

    for md_name, docx_name in files:
        md_path = os.path.join(INPUT_DIR, md_name)
        docx_path = os.path.join(OUTPUT_DIR, docx_name)
        if not os.path.exists(md_path):
            print(f"[跳过] 源文件不存在：{md_path}")
            continue
        try:
            convert_md_to_docx(md_path, docx_path)
            size_kb = os.path.getsize(docx_path) / 1024
            print(f"[成功] {docx_name}  ({size_kb:.1f} KB)")
        except Exception as e:
            print(f"[失败] {docx_name}: {e}")
            import traceback
            traceback.print_exc()

    print("-" * 60)
    print("转换完成。")


if __name__ == '__main__':
    main()
